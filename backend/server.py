from fastapi import FastAPI, APIRouter, HTTPException, Depends, status, File, UploadFile
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict, EmailStr
from typing import List, Optional
import uuid
from datetime import datetime, timezone, timedelta
from passlib.context import CryptContext
import jwt
from sendgrid import SendGridAPIClient
from sendgrid.helpers.mail import Mail
from apscheduler.schedulers.asyncio import AsyncIOScheduler
from apscheduler.triggers.cron import CronTrigger
import pytz
import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from io import BytesIO
import json
import base64
import shutil

ROOT_DIR = Path(__file__).parent
UPLOAD_DIR = ROOT_DIR / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Security
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
SECRET_KEY = os.environ.get('SECRET_KEY', 'your-secret-key-change-in-production')
ALGORITHM = "HS256"
security = HTTPBearer()

# Email configuration
SENDGRID_API_KEY = os.environ.get('SENDGRID_API_KEY', '')
ADMIN_EMAIL = os.environ.get('ADMIN_EMAIL', 'turyasin@gmail.com')

# Scheduler
scheduler = AsyncIOScheduler(timezone=pytz.timezone('Europe/Istanbul'))

# Helper functions for date calculations
def get_month_year(date_str: str) -> str:
    """Convert date string to 'Month YYYY' format in Turkish"""
    try:
        date_obj = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
        month_names = {
            1: "Ocak", 2: "Şubat", 3: "Mart", 4: "Nisan", 5: "Mayıs", 6: "Haziran",
            7: "Temmuz", 8: "Ağustos", 9: "Eylül", 10: "Ekim", 11: "Kasım", 12: "Aralık"
        }
        return f"{month_names[date_obj.month]} {date_obj.year}"
    except:
        return "N/A"

def get_quarter_year(date_str: str) -> str:
    """Convert date string to 'QX YYYY' format"""
    try:
        date_obj = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
        quarter = (date_obj.month - 1) // 3 + 1
        return f"Q{quarter} {date_obj.year}"
    except:
        return "N/A"

# Create the main app
app = FastAPI()
api_router = APIRouter(prefix="/api")

# Models
class User(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    username: str
    email: EmailStr
    is_admin: bool = False
    receive_notifications: bool = True
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class UserCreate(BaseModel):
    username: str
    email: EmailStr
    password: str

class UserUpdate(BaseModel):
    is_admin: Optional[bool] = None
    receive_notifications: Optional[bool] = None

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class Customer(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    email: Optional[str] = None
    phone: Optional[str] = None
    address: Optional[str] = None
    created_by: Optional[str] = None
    created_by_username: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class CustomerCreate(BaseModel):
    name: str
    email: Optional[str] = None
    phone: Optional[str] = None
    address: Optional[str] = None

class Invoice(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    customer_id: str
    customer_name: Optional[str] = None
    invoice_number: str
    amount: float
    paid_amount: float = 0.0
    due_date: str
    status: str = "unpaid"
    currency: str = "TRY"  # "TRY", "USD", "EUR"
    month: Optional[str] = None  # Auto-calculated: "Ocak 2025"
    quarter: Optional[str] = None  # Auto-calculated: "Q1 2025"
    notes: Optional[str] = None
    created_by: Optional[str] = None
    created_by_username: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class InvoiceCreate(BaseModel):
    customer_id: str
    invoice_number: str
    amount: float
    due_date: str
    currency: str = "TRY"  # "TRY", "USD", "EUR"
    notes: Optional[str] = None

class InvoiceUpdate(BaseModel):
    customer_id: Optional[str] = None
    invoice_number: Optional[str] = None
    amount: Optional[float] = None
    due_date: Optional[str] = None
    currency: Optional[str] = None  # "TRY", "USD", "EUR"
    notes: Optional[str] = None

class Payment(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    invoice_id: str
    invoice_number: Optional[str] = None
    customer_name: Optional[str] = None
    payment_method: str = "Çek"  # "Nakit", "Kredi Kartı", "Havale/EFT", "Çek"
    bank_account_id: Optional[str] = None  # For incoming payments
    bank_account_name: Optional[str] = None  # Populated from bank_account_id
    currency: Optional[str] = None  # Populated from bank_account_id or default TRY
    check_number: Optional[str] = None
    check_date: Optional[str] = None
    bank_name: Optional[str] = None
    amount: float
    month: Optional[str] = None  # Auto-calculated: "Ocak 2025"
    quarter: Optional[str] = None  # Auto-calculated: "Q1 2025"
    created_by: Optional[str] = None
    created_by_username: Optional[str] = None
    payment_date: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class PaymentCreate(BaseModel):
    invoice_id: str
    payment_method: str = "Çek"  # "Nakit", "Kredi Kartı", "Havale/EFT", "Çek"
    bank_account_id: Optional[str] = None  # For incoming payments
    check_number: Optional[str] = None
    check_date: Optional[str] = None
    bank_name: Optional[str] = None
    amount: float

class DashboardStats(BaseModel):
    total_invoices: int
    total_amount: float
    paid_amount: float
    outstanding_amount: float
    unpaid_count: int
    partial_count: int
    paid_count: int
    overdue_count: int  # Ödeme günü geçmiş faturalar
    recent_payments: List[Payment]
    total_received_checks: int
    total_received_amount: float
    pending_received_checks: int
    total_issued_checks: int
    total_issued_amount: float
    pending_issued_checks: int

# Çek Modeli - Alınan ve verilen çekler
class Check(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    check_type: str = "issued"  # "received" (alınan) veya "issued" (verilen) - default issued
    check_number: str
    amount: float
    due_date: str
    bank_name: str
    payer_payee: str  # Alıcı veya veren kişi/firma
    status: str = "pending"  # pending, collected, paid, bounced, cancelled
    related_invoice_id: Optional[str] = None  # Hangi fatura için kesildi (opsiyonel)
    notes: Optional[str] = None
    created_by: Optional[str] = None
    created_by_username: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class CheckCreate(BaseModel):
    check_type: str = "issued"  # Default issued (verilen çek)
    check_number: str
    amount: float
    due_date: str
    bank_name: str
    payer_payee: str
    related_invoice_id: Optional[str] = None
    notes: Optional[str] = None

class CheckUpdate(BaseModel):
    check_type: Optional[str] = None
    check_number: Optional[str] = None
    amount: Optional[float] = None
    due_date: Optional[str] = None
    bank_name: Optional[str] = None
    payer_payee: Optional[str] = None
    status: Optional[str] = None
    related_invoice_id: Optional[str] = None
    notes: Optional[str] = None

class WeeklyPaymentSchedule(BaseModel):
    week_label: str
    date_range: str
    received_checks: List[Check]
    issued_checks: List[Check]
    invoices_due: List[Invoice]
    total_receivable: float
    total_payable: float

class WeeklySchedule(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    week: str
    invoice_id: Optional[str] = None
    check_id: Optional[str] = None
    check_type: Optional[str]
    notes: Optional[str] = None

class BankAccount(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    bank_name: str
    branch: Optional[str] = None
    iban: str
    account_holder: str
    currency: str = "TRY"  # "TRY", "USD", "EUR"

class CompanyInfo(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    company_name: str
    tax_number: Optional[str] = None
    tax_office: Optional[str] = None
    address: Optional[str] = None
    phone: Optional[str] = None
    email: Optional[str] = None
    bank_accounts: list[BankAccount] = []
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    updated_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class CompanyInfoCreate(BaseModel):
    company_name: str
    tax_number: Optional[str] = None
    tax_office: Optional[str] = None
    address: Optional[str] = None
    phone: Optional[str] = None
    email: Optional[str] = None
    bank_accounts: list[BankAccount] = []

# Helper functions
def hash_password(password: str) -> str:
    return pwd_context.hash(password)

def verify_password(plain_password: str, hashed_password: str) -> bool:
    return pwd_context.verify(plain_password, hashed_password)

def create_access_token(data: dict) -> str:
    to_encode = data.copy()
    expire = datetime.now(timezone.utc) + timedelta(days=7)
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        token = credentials.credentials
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        user_id = payload.get("sub")
        if user_id is None:
            raise HTTPException(status_code=401, detail="Invalid token")
        return user_id
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

async def get_current_admin_user(user_id: str = Depends(get_current_user)):
    """Check if current user is admin"""
    user = await db.users.find_one({"id": user_id}, {"_id": 0})
    if not user or not user.get("is_admin", False):
        raise HTTPException(status_code=403, detail="Admin yetkisi gerekli")
    return user_id

# Email notification functions
async def send_invoice_reminder_email(invoice_number: str, customer_name: str, amount: float, due_date: str, recipient_email: str, recipient_name: str):
    """Send email reminder for upcoming invoice due date"""
    if not SENDGRID_API_KEY or SENDGRID_API_KEY == 'your-sendgrid-api-key-here':
        logger.warning(f"SendGrid API key not configured. Skipping email for invoice {invoice_number}")
        return False
    
    try:
        message = Mail(
            from_email='noreply@odemetakip.com',
            to_emails=recipient_email,
            subject=f'Ödeme Hatırlatması: Fatura {invoice_number} - 2 Gün Sonra',
            html_content=f'''
            <html>
            <body style="font-family: Arial, sans-serif; padding: 20px;">
                <h2 style="color: #2563eb;">Ödeme Hatırlatması</h2>
                <p>Sayın {recipient_name},</p>
                <p>Aşağıdaki faturanın vadesi 2 gün sonra dolmaktadır:</p>
                <div style="background-color: #f1f5f9; padding: 15px; border-radius: 8px; margin: 20px 0;">
                    <p><strong>Fatura No:</strong> {invoice_number}</p>
                    <p><strong>Müşteri:</strong> {customer_name}</p>
                    <p><strong>Tutar:</strong> ₺{amount:.2f}</p>
                    <p><strong>Vade Tarihi:</strong> {due_date}</p>
                </div>
                <p style="color: #dc2626;">Lütfen ödeme tahsilatının planlandığından emin olun.</p>
                <hr style="margin: 20px 0;">
                <p style="color: #64748b; font-size: 12px;">Bu otomatik bir hatırlatmadır - Ödeme Takip Sistemi</p>
            </body>
            </html>
            '''
        )
        
        sg = SendGridAPIClient(SENDGRID_API_KEY)
        response = sg.send(message)
        logger.info(f"Email sent to {recipient_email} for invoice {invoice_number}. Status: {response.status_code}")
        return True
    except Exception as e:
        logger.error(f"Failed to send email to {recipient_email} for invoice {invoice_number}: {str(e)}")
        return False

async def check_upcoming_invoices():
    """Check for invoices due in 2 days and send reminders"""
    try:
        logger.info("Running daily invoice reminder check...")
        
        # Calculate date 2 days from now
        target_date = (datetime.now(timezone.utc) + timedelta(days=2)).date()
        target_date_str = target_date.isoformat()
        
        # Find unpaid and partial invoices due on target date
        invoices = await db.invoices.find({
            "status": {"$in": ["unpaid", "partial"]},
            "due_date": target_date_str
        }, {"_id": 0}).to_list(1000)
        
        if not invoices:
            logger.info(f"No invoices due on {target_date_str}")
            return
        
        logger.info(f"Found {len(invoices)} invoices due on {target_date_str}")
        
        # Get users who want to receive notifications
        notification_users = await db.users.find({
            "receive_notifications": True
        }, {"_id": 0, "email": 1, "username": 1}).to_list(1000)
        
        if not notification_users:
            logger.info("No users configured to receive notifications")
            return
        
        # Send email for each invoice to all notification users
        for invoice in invoices:
            # Get customer name
            customer = await db.customers.find_one({"id": invoice["customer_id"]}, {"_id": 0})
            customer_name = customer["name"] if customer else "Unknown Customer"
            
            # Send to all notification users
            for user in notification_users:
                await send_invoice_reminder_email(
                    invoice_number=invoice["invoice_number"],
                    customer_name=customer_name,
                    amount=invoice["amount"],
                    due_date=invoice["due_date"],
                    recipient_email=user["email"],
                    recipient_name=user["username"]
                )
        
        logger.info("Daily invoice reminder check completed")
    except Exception as e:
        logger.error(f"Error in check_upcoming_invoices: {str(e)}")

# Auth routes
@api_router.post("/auth/register", response_model=dict)
async def register(user: UserCreate):
    existing = await db.users.find_one({"email": user.email})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    # Check if this is the first user (make them admin)
    user_count = await db.users.count_documents({})
    
    user_dict = user.model_dump()
    user_dict["password"] = hash_password(user_dict["password"])
    user_obj = User(**{k: v for k, v in user_dict.items() if k != "password"})
    
    # First user OR Yasin's email becomes admin
    if user_count == 0 or user.email.lower() == "turyasin@gmail.com":
        user_obj.is_admin = True
    
    doc = user_obj.model_dump()
    doc["password"] = user_dict["password"]
    await db.users.insert_one(doc)
    
    token = create_access_token({"sub": user_obj.id})
    return {"token": token, "user": user_obj.model_dump()}

@api_router.post("/auth/login", response_model=dict)
async def login(credentials: UserLogin):
    user = await db.users.find_one({"email": credentials.email})
    if not user or not verify_password(credentials.password, user["password"]):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    token = create_access_token({"sub": user["id"]})
    user_data = User(**user).model_dump()
    return {"token": token, "user": user_data}

# User routes
@api_router.get("/users", response_model=List[User])
async def get_users(user_id: str = Depends(get_current_user)):
    users = await db.users.find({}, {"_id": 0, "password": 0}).to_list(1000)
    return users

@api_router.get("/users/me", response_model=User)
async def get_current_user_info(user_id: str = Depends(get_current_user)):
    user = await db.users.find_one({"id": user_id}, {"_id": 0, "password": 0})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    return user

@api_router.put("/users/{target_user_id}", response_model=User)
async def update_user(target_user_id: str, user_update: UserUpdate, admin_id: str = Depends(get_current_admin_user)):
    """Update user permissions (admin only)"""
    target_user = await db.users.find_one({"id": target_user_id}, {"_id": 0})
    if not target_user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Don't allow admin to remove their own admin status if they're the only admin
    if target_user_id == admin_id and user_update.is_admin is False:
        admin_count = await db.users.count_documents({"is_admin": True})
        if admin_count <= 1:
            raise HTTPException(status_code=400, detail="En az bir admin olmalı")
    
    update_data = user_update.model_dump(exclude_unset=True)
    await db.users.update_one({"id": target_user_id}, {"$set": update_data})
    
    updated_user = await db.users.find_one({"id": target_user_id}, {"_id": 0, "password": 0})
    return updated_user

@api_router.delete("/users/{target_user_id}")
async def delete_user(target_user_id: str, admin_id: str = Depends(get_current_admin_user)):
    """Delete user (admin only)"""
    # Check if user exists
    target_user = await db.users.find_one({"id": target_user_id}, {"_id": 0})
    if not target_user:
        raise HTTPException(status_code=404, detail="Kullanıcı bulunamadı")
    
    # Don't allow admin to delete themselves
    if target_user_id == admin_id:
        raise HTTPException(status_code=400, detail="Kendi hesabınızı silemezsiniz")
    
    # Don't allow deleting the last admin
    if target_user.get("is_admin"):
        admin_count = await db.users.count_documents({"is_admin": True})
        if admin_count <= 1:
            raise HTTPException(status_code=400, detail="En az bir admin olmalı")
    
    # Delete the user
    result = await db.users.delete_one({"id": target_user_id})
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=500, detail="Kullanıcı silinemedi")
    
    return {"message": f"Kullanıcı {target_user.get('username')} başarıyla silindi", "deleted_user_id": target_user_id}

# Customer routes
@api_router.get("/customers", response_model=List[Customer])
async def get_customers(user_id: str = Depends(get_current_user)):
    customers = await db.customers.find({}, {"_id": 0}).to_list(1000)
    return customers

@api_router.post("/customers", response_model=Customer)
async def create_customer(customer: CustomerCreate, user_id: str = Depends(get_current_user)):
    # Get current user info
    current_user = await db.users.find_one({"id": user_id}, {"_id": 0})
    
    customer_obj = Customer(**customer.model_dump())
    customer_obj.created_by = user_id
    customer_obj.created_by_username = current_user.get("username", "Unknown") if current_user else "Unknown"
    
    await db.customers.insert_one(customer_obj.model_dump())
    return customer_obj

@api_router.get("/customers/{customer_id}", response_model=Customer)
async def get_customer(customer_id: str, user_id: str = Depends(get_current_user)):
    customer = await db.customers.find_one({"id": customer_id}, {"_id": 0})
    if not customer:
        raise HTTPException(status_code=404, detail="Customer not found")
    return customer

@api_router.put("/customers/{customer_id}", response_model=Customer)
async def update_customer(customer_id: str, customer: CustomerCreate, user_id: str = Depends(get_current_user)):
    result = await db.customers.find_one({"id": customer_id}, {"_id": 0})
    if not result:
        raise HTTPException(status_code=404, detail="Customer not found")
    
    await db.customers.update_one({"id": customer_id}, {"$set": customer.model_dump(exclude_unset=True)})
    updated = await db.customers.find_one({"id": customer_id}, {"_id": 0})
    return updated

@api_router.delete("/customers/{customer_id}")
async def delete_customer(customer_id: str, user_id: str = Depends(get_current_user)):
    result = await db.customers.delete_one({"id": customer_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Customer not found")
    return {"message": "Customer deleted"}

# Invoice routes
@api_router.get("/invoices", response_model=List[Invoice])
async def get_invoices(status: Optional[str] = None, user_id: str = Depends(get_current_user)):
    query = {}
    if status:
        query["status"] = status
    
    invoices = await db.invoices.find(query, {"_id": 0}).to_list(1000)
    
    for invoice in invoices:
        customer = await db.customers.find_one({"id": invoice["customer_id"]}, {"_id": 0})
        if customer:
            invoice["customer_name"] = customer["name"]
    
    return invoices

@api_router.post("/invoices", response_model=Invoice)
async def create_invoice(invoice: InvoiceCreate, user_id: str = Depends(get_current_user)):
    existing = await db.invoices.find_one({"invoice_number": invoice.invoice_number})
    if existing:
        raise HTTPException(status_code=400, detail="Invoice number already exists")
    
    # Get current user info
    current_user = await db.users.find_one({"id": user_id}, {"_id": 0})
    
    invoice_obj = Invoice(**invoice.model_dump())
    invoice_obj.created_by = user_id
    invoice_obj.created_by_username = current_user.get("username", "Unknown") if current_user else "Unknown"
    
    # Auto-calculate month and quarter from due_date
    invoice_obj.month = get_month_year(invoice.due_date)
    invoice_obj.quarter = get_quarter_year(invoice.due_date)
    
    customer = await db.customers.find_one({"id": invoice.customer_id}, {"_id": 0})
    if customer:
        invoice_obj.customer_name = customer["name"]
    
    await db.invoices.insert_one(invoice_obj.model_dump())
    return invoice_obj

@api_router.get("/invoices/{invoice_id}", response_model=Invoice)
async def get_invoice(invoice_id: str, user_id: str = Depends(get_current_user)):
    invoice = await db.invoices.find_one({"id": invoice_id}, {"_id": 0})
    if not invoice:
        raise HTTPException(status_code=404, detail="Invoice not found")
    
    customer = await db.customers.find_one({"id": invoice["customer_id"]}, {"_id": 0})
    if customer:
        invoice["customer_name"] = customer["name"]
    
    return invoice

@api_router.put("/invoices/{invoice_id}", response_model=Invoice)
async def update_invoice(invoice_id: str, invoice: InvoiceUpdate, user_id: str = Depends(get_current_user)):
    result = await db.invoices.find_one({"id": invoice_id}, {"_id": 0})
    if not result:
        raise HTTPException(status_code=404, detail="Invoice not found")
    
    update_data = invoice.model_dump(exclude_unset=True)
    
    # If due_date is updated, recalculate month and quarter
    if "due_date" in update_data:
        update_data["month"] = get_month_year(update_data["due_date"])
        update_data["quarter"] = get_quarter_year(update_data["due_date"])
    
    await db.invoices.update_one({"id": invoice_id}, {"$set": update_data})
    
    updated = await db.invoices.find_one({"id": invoice_id}, {"_id": 0})
    
    customer = await db.customers.find_one({"id": updated["customer_id"]}, {"_id": 0})
    if customer:
        updated["customer_name"] = customer["name"]
    
    return updated

@api_router.delete("/invoices/{invoice_id}")
async def delete_invoice(invoice_id: str, user_id: str = Depends(get_current_user)):
    result = await db.invoices.delete_one({"id": invoice_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Invoice not found")
    
    await db.payments.delete_many({"invoice_id": invoice_id})
    return {"message": "Invoice deleted"}

# Payment routes
@api_router.get("/payments", response_model=List[Payment])
async def get_payments(user_id: str = Depends(get_current_user)):
    payments = await db.payments.find({}, {"_id": 0}).to_list(1000)
    
    for payment in payments:
        invoice = await db.invoices.find_one({"id": payment["invoice_id"]}, {"_id": 0})
        if invoice:
            payment["invoice_number"] = invoice["invoice_number"]
            customer = await db.customers.find_one({"id": invoice["customer_id"]}, {"_id": 0})
            if customer:
                payment["customer_name"] = customer["name"]
    
    return payments

@api_router.post("/payments", response_model=Payment)
async def create_payment(payment: PaymentCreate, user_id: str = Depends(get_current_user)):
    invoice = await db.invoices.find_one({"id": payment.invoice_id}, {"_id": 0})
    if not invoice:
        raise HTTPException(status_code=404, detail="Invoice not found")
    
    # Get current user info
    current_user = await db.users.find_one({"id": user_id}, {"_id": 0})
    
    payment_obj = Payment(**payment.model_dump())
    payment_obj.invoice_number = invoice["invoice_number"]
    payment_obj.created_by = user_id
    payment_obj.created_by_username = current_user.get("username", "Unknown") if current_user else "Unknown"
    
    # Auto-calculate month and quarter from payment_date
    payment_obj.month = get_month_year(payment_obj.payment_date)
    payment_obj.quarter = get_quarter_year(payment_obj.payment_date)
    
    # Populate bank account info if bank_account_id is provided
    if payment_obj.bank_account_id:
        company_info = await db.company_info.find_one({}, {"_id": 0})
        if company_info:
            for account in company_info.get("bank_accounts", []):
                if account.get("id") == payment_obj.bank_account_id:
                    payment_obj.bank_account_name = account.get("bank_name")
                    payment_obj.currency = account.get("currency", "TRY")
                    break
    
    # Default currency if not set
    if not payment_obj.currency:
        payment_obj.currency = "TRY"
    
    customer = await db.customers.find_one({"id": invoice["customer_id"]}, {"_id": 0})
    if customer:
        payment_obj.customer_name = customer["name"]
    
    await db.payments.insert_one(payment_obj.model_dump())
    
    new_paid_amount = invoice.get("paid_amount", 0) + payment.amount
    invoice_amount = invoice["amount"]
    
    if new_paid_amount >= invoice_amount:
        new_status = "paid"
    elif new_paid_amount > 0:
        new_status = "partial"
    else:
        new_status = "unpaid"
    
    await db.invoices.update_one(
        {"id": payment.invoice_id},
        {"$set": {"paid_amount": new_paid_amount, "status": new_status}}
    )
    
    return payment_obj

@api_router.get("/payments/invoice/{invoice_id}", response_model=List[Payment])
async def get_invoice_payments(invoice_id: str, user_id: str = Depends(get_current_user)):
    payments = await db.payments.find({"invoice_id": invoice_id}, {"_id": 0}).to_list(1000)
    return payments

@api_router.delete("/payments/{payment_id}")
async def delete_payment(payment_id: str, user_id: str = Depends(get_current_user)):
    payment = await db.payments.find_one({"id": payment_id}, {"_id": 0})
    if not payment:
        raise HTTPException(status_code=404, detail="Payment not found")
    
    invoice = await db.invoices.find_one({"id": payment["invoice_id"]}, {"_id": 0})
    if invoice:
        new_paid_amount = max(0, invoice.get("paid_amount", 0) - payment["amount"])
        invoice_amount = invoice["amount"]
        
        if new_paid_amount >= invoice_amount:
            new_status = "paid"
        elif new_paid_amount > 0:
            new_status = "partial"
        else:
            new_status = "unpaid"
        
        await db.invoices.update_one(
            {"id": payment["invoice_id"]},
            {"$set": {"paid_amount": new_paid_amount, "status": new_status}}
        )
    
    await db.payments.delete_one({"id": payment_id})
    return {"message": "Payment deleted"}

# Dashboard route
@api_router.get("/dashboard/stats", response_model=DashboardStats)
async def get_dashboard_stats(user_id: str = Depends(get_current_user)):
    invoices = await db.invoices.find({}, {"_id": 0}).to_list(1000)
    payments = await db.payments.find({}, {"_id": 0}).sort("created_at", -1).to_list(5)
    
    total_invoices = len(invoices)
    total_amount = sum(inv["amount"] for inv in invoices)
    paid_amount = sum(inv.get("paid_amount", 0) for inv in invoices)
    outstanding_amount = total_amount - paid_amount
    
    unpaid_count = len([inv for inv in invoices if inv["status"] == "unpaid"])
    partial_count = len([inv for inv in invoices if inv["status"] == "partial"])
    paid_count = len([inv for inv in invoices if inv["status"] == "paid"])
    
    # Calculate overdue invoices (due date has passed and not fully paid)
    today = datetime.now(timezone.utc).date()
    overdue_count = len([
        inv for inv in invoices 
        if inv["status"] in ["unpaid", "partial"] and 
        datetime.fromisoformat(inv["due_date"].replace('Z', '+00:00')).date() < today
    ])
    
    for payment in payments:
        invoice = await db.invoices.find_one({"id": payment["invoice_id"]}, {"_id": 0})
        if invoice:
            payment["invoice_number"] = invoice["invoice_number"]
            customer = await db.customers.find_one({"id": invoice["customer_id"]}, {"_id": 0})
            if customer:
                payment["customer_name"] = customer["name"]
    
    # Check statistics
    received_checks = await db.checks.find({"check_type": "received"}, {"_id": 0}).to_list(1000)
    issued_checks = await db.checks.find({"check_type": "issued"}, {"_id": 0}).to_list(1000)
    
    total_received_checks = len(received_checks)
    total_received_amount = sum(c["amount"] for c in received_checks)
    pending_received_checks = len([c for c in received_checks if c["status"] == "pending"])
    
    total_issued_checks = len(issued_checks)
    total_issued_amount = sum(c["amount"] for c in issued_checks)
    pending_issued_checks = len([c for c in issued_checks if c["status"] == "pending"])
    
    return DashboardStats(
        total_invoices=total_invoices,
        total_amount=total_amount,
        paid_amount=paid_amount,
        outstanding_amount=outstanding_amount,
        unpaid_count=unpaid_count,
        partial_count=partial_count,
        paid_count=paid_count,
        recent_payments=payments,
        total_received_checks=total_received_checks,
        total_received_amount=total_received_amount,
        pending_received_checks=pending_received_checks,
        total_issued_checks=total_issued_checks,
        total_issued_amount=total_issued_amount,
        pending_issued_checks=pending_issued_checks
    )

# Email notification endpoint (for testing)
@api_router.post("/notifications/check-reminders")
async def trigger_reminder_check(user_id: str = Depends(get_current_user)):
    """Manually trigger invoice reminder check"""
    await check_upcoming_invoices()
    return {"message": "Invoice reminder check completed"}

# Check routes
@api_router.get("/checks", response_model=List[Check])
async def get_checks(check_type: Optional[str] = None, status: Optional[str] = None, user_id: str = Depends(get_current_user)):
    query = {}
    if check_type:
        query["check_type"] = check_type
    if status:
        query["status"] = status
    
    checks = await db.checks.find(query, {"_id": 0}).to_list(1000)
    return checks

@api_router.post("/checks", response_model=Check)
async def create_check(check: CheckCreate, user_id: str = Depends(get_current_user)):
    # Get current user info
    current_user = await db.users.find_one({"id": user_id}, {"_id": 0})
    
    check_obj = Check(**check.model_dump())
    check_obj.created_by = user_id
    check_obj.created_by_username = current_user.get("username", "Unknown") if current_user else "Unknown"
    
    await db.checks.insert_one(check_obj.model_dump())
    return check_obj

@api_router.get("/checks/{check_id}", response_model=Check)
async def get_check(check_id: str, user_id: str = Depends(get_current_user)):
    check = await db.checks.find_one({"id": check_id}, {"_id": 0})
    if not check:
        raise HTTPException(status_code=404, detail="Check not found")
    return check

@api_router.put("/checks/{check_id}", response_model=Check)
async def update_check(check_id: str, check: CheckUpdate, user_id: str = Depends(get_current_user)):
    result = await db.checks.find_one({"id": check_id}, {"_id": 0})
    if not result:
        raise HTTPException(status_code=404, detail="Check not found")
    
    update_data = check.model_dump(exclude_unset=True)
    await db.checks.update_one({"id": check_id}, {"$set": update_data})
    
    updated = await db.checks.find_one({"id": check_id}, {"_id": 0})
    return updated

@api_router.delete("/checks/{check_id}")
async def delete_check(check_id: str, user_id: str = Depends(get_current_user)):
    result = await db.checks.delete_one({"id": check_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Check not found")
    return {"message": "Check deleted"}

# Weekly payment schedule
@api_router.get("/payments/weekly-schedule", response_model=List[WeeklyPaymentSchedule])
async def get_weekly_payment_schedule(weeks: int = 4, user_id: str = Depends(get_current_user)):
    """Get weekly payment schedule for next N weeks"""
    schedule = []
    today = datetime.now(timezone.utc).date()
    
    for week_num in range(weeks):
        week_start = today + timedelta(weeks=week_num)
        week_end = week_start + timedelta(days=6)
        
        week_start_str = week_start.isoformat()
        week_end_str = week_end.isoformat()
        
        # Get received checks due this week
        received_checks = await db.checks.find({
            "check_type": "received",
            "status": {"$in": ["pending"]},
            "due_date": {"$gte": week_start_str, "$lte": week_end_str}
        }, {"_id": 0}).to_list(1000)
        
        # Get issued checks due this week
        issued_checks = await db.checks.find({
            "check_type": "issued",
            "status": {"$in": ["pending"]},
            "due_date": {"$gte": week_start_str, "$lte": week_end_str}
        }, {"_id": 0}).to_list(1000)
        
        # Get invoices due this week
        invoices_due = await db.invoices.find({
            "status": {"$in": ["unpaid", "partial"]},
            "due_date": {"$gte": week_start_str, "$lte": week_end_str}
        }, {"_id": 0}).to_list(1000)
        
        # Enrich invoices with customer names
        for invoice in invoices_due:
            customer = await db.customers.find_one({"id": invoice["customer_id"]}, {"_id": 0})
            if customer:
                invoice["customer_name"] = customer["name"]
        
        total_receivable = sum(c["amount"] for c in received_checks) + sum(i["amount"] - i.get("paid_amount", 0) for i in invoices_due)
        total_payable = sum(c["amount"] for c in issued_checks)
        
        week_label = "Bu Hafta" if week_num == 0 else f"{week_num + 1}. Hafta"
        date_range = f"{week_start.strftime('%d.%m')} - {week_end.strftime('%d.%m.%Y')}"
        
        schedule.append(WeeklyPaymentSchedule(
            week_label=week_label,
            date_range=date_range,
            received_checks=received_checks,
            issued_checks=issued_checks,
            invoices_due=invoices_due,
            total_receivable=total_receivable,
            total_payable=total_payable
        ))
    
    return schedule

# Helper functions for export
async def export_invoices_xlsx(invoices):
    """Export invoices to Excel format"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Faturalar"
    
    # Headers
    headers = ["ID", "Müşteri ID", "Müşteri Adı", "Fatura No", "Tutar (₺)", "Ödenen (₺)", "Vade Tarihi", "Durum", "Periyot", "Notlar", "Oluşturan", "Oluşturma Tarihi"]
    ws.append(headers)
    
    # Style headers
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        cell.alignment = Alignment(horizontal="center")
    
    # Data rows
    for inv in invoices:
        ws.append([
            inv.get("id", ""),
            inv.get("customer_id", ""),
            inv.get("customer_name", ""),
            inv.get("invoice_number", ""),
            inv.get("amount", 0),
            inv.get("paid_amount", 0),
            inv.get("due_date", ""),
            inv.get("status", ""),
            inv.get("period_type", "Aylık"),
            inv.get("notes", ""),
            inv.get("created_by_username", ""),
            inv.get("created_at", "")
        ])
    
    # Auto-adjust column widths
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width
    
    # Save to BytesIO
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename=faturalar_{datetime.now().strftime('%Y%m%d')}.xlsx"}
    )

async def export_invoices_docx(invoices):
    """Export invoices to Word format"""
    doc = Document()
    doc.add_heading('Faturalar Raporu', 0)
    doc.add_paragraph(f'Rapor Tarihi: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    doc.add_paragraph(f'Toplam Fatura: {len(invoices)}')
    doc.add_paragraph('')
    
    for inv in invoices:
        doc.add_heading(f'Fatura No: {inv.get("invoice_number", "")}', level=2)
        doc.add_paragraph(f'Müşteri: {inv.get("customer_name", "")}')
        doc.add_paragraph(f'Tutar: ₺{inv.get("amount", 0):,.2f}')
        doc.add_paragraph(f'Ödenen: ₺{inv.get("paid_amount", 0):,.2f}')
        doc.add_paragraph(f'Vade Tarihi: {inv.get("due_date", "")}')
        doc.add_paragraph(f'Durum: {inv.get("status", "")}')
        doc.add_paragraph(f'Periyot: {inv.get("period_type", "Aylık")}')
        if inv.get("notes"):
            doc.add_paragraph(f'Notlar: {inv.get("notes", "")}')
        doc.add_paragraph(f'Oluşturan: {inv.get("created_by_username", "")}')
        doc.add_paragraph('')
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=faturalar_{datetime.now().strftime('%Y%m%d')}.docx"}
    )

async def export_invoices_pdf(invoices):
    """Export invoices to PDF format"""
    output = BytesIO()
    doc = SimpleDocTemplate(output, pagesize=landscape(A4))
    elements = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=20, textColor=colors.HexColor('#366092'))
    elements.append(Paragraph('Faturalar Raporu', title_style))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph(f'Rapor Tarihi: {datetime.now().strftime("%d.%m.%Y %H:%M")}', styles['Normal']))
    elements.append(Paragraph(f'Toplam Fatura: {len(invoices)}', styles['Normal']))
    elements.append(Spacer(1, 20))
    
    # Table data
    data = [['Fatura No', 'Müşteri', 'Tutar (₺)', 'Ödenen (₺)', 'Vade', 'Durum', 'Periyot', 'Oluşturan']]
    for inv in invoices:
        data.append([
            inv.get("invoice_number", ""),
            inv.get("customer_name", "")[:20],
            f"₺{inv.get('amount', 0):,.2f}",
            f"₺{inv.get('paid_amount', 0):,.2f}",
            inv.get("due_date", ""),
            inv.get("status", ""),
            inv.get("period_type", "Aylık"),
            inv.get("created_by_username", "")[:15]
        ])
    
    table = Table(data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#366092')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    elements.append(table)
    doc.build(elements)
    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=faturalar_{datetime.now().strftime('%Y%m%d')}.pdf"}
    )

async def export_checks_xlsx(checks):
    """Export checks to Excel format"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Çekler"
    
    # Headers
    headers = ["ID", "Tür", "Çek No", "Tutar (₺)", "Vade Tarihi", "Banka", "Alıcı/Veren", "Durum", "Notlar", "Oluşturan", "Oluşturma Tarihi"]
    ws.append(headers)
    
    # Style headers
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        cell.alignment = Alignment(horizontal="center")
    
    # Data rows
    for check in checks:
        check_type_tr = "Alınan" if check.get("check_type") == "received" else "Verilen"
        ws.append([
            check.get("id", ""),
            check_type_tr,
            check.get("check_number", ""),
            check.get("amount", 0),
            check.get("due_date", ""),
            check.get("bank_name", ""),
            check.get("payer_payee", ""),
            check.get("status", ""),
            check.get("notes", ""),
            check.get("created_by_username", ""),
            check.get("created_at", "")
        ])
    
    # Auto-adjust column widths
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width
    
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename=cekler_{datetime.now().strftime('%Y%m%d')}.xlsx"}
    )

async def export_checks_docx(checks):
    """Export checks to Word format"""
    doc = Document()
    doc.add_heading('Çekler Raporu', 0)
    doc.add_paragraph(f'Rapor Tarihi: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    doc.add_paragraph(f'Toplam Çek: {len(checks)}')
    doc.add_paragraph('')
    
    for check in checks:
        check_type_tr = "Alınan Çek" if check.get("check_type") == "received" else "Verilen Çek"
        doc.add_heading(f'{check_type_tr} - {check.get("check_number", "")}', level=2)
        doc.add_paragraph(f'Tutar: ₺{check.get("amount", 0):,.2f}')
        doc.add_paragraph(f'Vade Tarihi: {check.get("due_date", "")}')
        doc.add_paragraph(f'Banka: {check.get("bank_name", "")}')
        doc.add_paragraph(f'Alıcı/Veren: {check.get("payer_payee", "")}')
        doc.add_paragraph(f'Durum: {check.get("status", "")}')
        if check.get("notes"):
            doc.add_paragraph(f'Notlar: {check.get("notes", "")}')
        doc.add_paragraph(f'Oluşturan: {check.get("created_by_username", "")}')
        doc.add_paragraph('')
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=cekler_{datetime.now().strftime('%Y%m%d')}.docx"}
    )

async def export_checks_pdf(checks):
    """Export checks to PDF format"""
    output = BytesIO()
    doc = SimpleDocTemplate(output, pagesize=landscape(A4))
    elements = []
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=20, textColor=colors.HexColor('#366092'))
    elements.append(Paragraph('Çekler Raporu', title_style))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph(f'Rapor Tarihi: {datetime.now().strftime("%d.%m.%Y %H:%M")}', styles['Normal']))
    elements.append(Paragraph(f'Toplam Çek: {len(checks)}', styles['Normal']))
    elements.append(Spacer(1, 20))
    
    data = [['Tür', 'Çek No', 'Tutar (₺)', 'Vade', 'Banka', 'Alıcı/Veren', 'Durum', 'Oluşturan']]
    for check in checks:
        check_type_tr = "Alınan" if check.get("check_type") == "received" else "Verilen"
        data.append([
            check_type_tr,
            check.get("check_number", ""),
            f"₺{check.get('amount', 0):,.2f}",
            check.get("due_date", ""),
            check.get("bank_name", "")[:15],
            check.get("payer_payee", "")[:15],
            check.get("status", ""),
            check.get("created_by_username", "")[:12]
        ])
    
    table = Table(data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#366092')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    elements.append(table)
    doc.build(elements)
    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=cekler_{datetime.now().strftime('%Y%m%d')}.pdf"}
    )

async def export_payments_xlsx(payments):
    """Export payments to Excel format"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Ödemeler"
    
    headers = ["ID", "Fatura ID", "Fatura No", "Müşteri", "Çek No", "Çek Tarihi", "Banka", "Tutar (₺)", "Periyot", "Oluşturan", "Ödeme Tarihi"]
    ws.append(headers)
    
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        cell.alignment = Alignment(horizontal="center")
    
    for payment in payments:
        ws.append([
            payment.get("id", ""),
            payment.get("invoice_id", ""),
            payment.get("invoice_number", ""),
            payment.get("customer_name", ""),
            payment.get("check_number", ""),
            payment.get("check_date", ""),
            payment.get("bank_name", ""),
            payment.get("amount", 0),
            payment.get("period_type", "Aylık"),
            payment.get("created_by_username", ""),
            payment.get("payment_date", "")
        ])
    
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width
    
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename=odemeler_{datetime.now().strftime('%Y%m%d')}.xlsx"}
    )

async def export_payments_docx(payments):
    """Export payments to Word format"""
    doc = Document()
    doc.add_heading('Ödemeler Raporu', 0)
    doc.add_paragraph(f'Rapor Tarihi: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    doc.add_paragraph(f'Toplam Ödeme: {len(payments)}')
    doc.add_paragraph('')
    
    for payment in payments:
        doc.add_heading(f'Ödeme - {payment.get("check_number", "")}', level=2)
        doc.add_paragraph(f'Fatura No: {payment.get("invoice_number", "")}')
        doc.add_paragraph(f'Müşteri: {payment.get("customer_name", "")}')
        doc.add_paragraph(f'Tutar: ₺{payment.get("amount", 0):,.2f}')
        doc.add_paragraph(f'Çek Tarihi: {payment.get("check_date", "")}')
        doc.add_paragraph(f'Banka: {payment.get("bank_name", "")}')
        doc.add_paragraph(f'Periyot: {payment.get("period_type", "Aylık")}')
        doc.add_paragraph(f'Oluşturan: {payment.get("created_by_username", "")}')
        doc.add_paragraph('')
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=odemeler_{datetime.now().strftime('%Y%m%d')}.docx"}
    )

async def export_payments_pdf(payments):
    """Export payments to PDF format"""
    output = BytesIO()
    doc = SimpleDocTemplate(output, pagesize=landscape(A4))
    elements = []
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=20, textColor=colors.HexColor('#366092'))
    elements.append(Paragraph('Ödemeler Raporu', title_style))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph(f'Rapor Tarihi: {datetime.now().strftime("%d.%m.%Y %H:%M")}', styles['Normal']))
    elements.append(Paragraph(f'Toplam Ödeme: {len(payments)}', styles['Normal']))
    elements.append(Spacer(1, 20))
    
    data = [['Fatura No', 'Müşteri', 'Çek No', 'Çek Tarihi', 'Banka', 'Tutar (₺)', 'Periyot', 'Oluşturan']]
    for payment in payments:
        data.append([
            payment.get("invoice_number", "")[:15],
            payment.get("customer_name", "")[:20],
            payment.get("check_number", ""),
            payment.get("check_date", ""),
            payment.get("bank_name", "")[:15],
            f"₺{payment.get('amount', 0):,.2f}",
            payment.get("period_type", "Aylık"),
            payment.get("created_by_username", "")[:12]
        ])
    
    table = Table(data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#366092')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    elements.append(table)
    doc.build(elements)
    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=odemeler_{datetime.now().strftime('%Y%m%d')}.pdf"}
    )

async def export_weekly_schedule_xlsx(schedule):
    """Export weekly schedule to Excel format"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Haftalık Plan"
    
    ws.append(["Haftalık Ödeme Planı"])
    ws.merge_cells('A1:H1')
    ws['A1'].font = Font(bold=True, size=16)
    ws['A1'].alignment = Alignment(horizontal="center")
    ws.append([])
    
    for week in schedule:
        ws.append([f"{week.week_label} ({week.date_range})"])
        ws.merge_cells(f'A{ws.max_row}:H{ws.max_row}')
        ws[f'A{ws.max_row}'].font = Font(bold=True, color="FFFFFF")
        ws[f'A{ws.max_row}'].fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        
        ws.append(["Alınan Çekler", "", "", "", "", "", "", ""])
        headers_row = ws.max_row
        ws.append(["Çek No", "Tutar (₺)", "Vade", "Banka", "Veren", "Durum", "", ""])
        for check in week.received_checks:
            ws.append([
                check.check_number,
                check.amount,
                check.due_date,
                check.bank_name,
                check.payer_payee,
                check.status,
                "",
                ""
            ])
        
        ws.append(["Verilen Çekler", "", "", "", "", "", "", ""])
        ws.append(["Çek No", "Tutar (₺)", "Vade", "Banka", "Alan", "Durum", "", ""])
        for check in week.issued_checks:
            ws.append([
                check.check_number,
                check.amount,
                check.due_date,
                check.bank_name,
                check.payer_payee,
                check.status,
                "",
                ""
            ])
        
        ws.append(["Vadesi Gelen Faturalar", "", "", "", "", "", "", ""])
        ws.append(["Fatura No", "Müşteri", "Tutar (₺)", "Vade", "", "", "", ""])
        for invoice in week.invoices_due:
            ws.append([
                invoice.invoice_number,
                invoice.customer_name,
                invoice.amount,
                invoice.due_date,
                "",
                "",
                "",
                ""
            ])
        
        ws.append(["Toplam Alacak:", f"₺{week.total_receivable:,.2f}", "Toplam Borç:", f"₺{week.total_payable:,.2f}", "", "", "", ""])
        ws[f'A{ws.max_row}'].font = Font(bold=True)
        ws[f'C{ws.max_row}'].font = Font(bold=True)
        ws.append([])
    
#     for column in ws.columns:
#         max_length = 0
#         column_letter = column[0].column_letter
#         for cell in column:
#             try:
#                 if len(str(cell.value)) > max_length:
#                     max_length = len(str(cell.value))
#             except:
#                 pass
#         adjusted_width = min(max_length + 2, 50)
#         ws.column_dimensions[column_letter].width = adjusted_width
    
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename=haftalik_plan_{datetime.now().strftime('%Y%m%d')}.xlsx"}
    )

async def export_weekly_schedule_docx(schedule):
    """Export weekly schedule to Word format"""
    doc = Document()
    doc.add_heading('Haftalık Ödeme Planı', 0)
    doc.add_paragraph(f'Rapor Tarihi: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    doc.add_paragraph('')
    
    for week in schedule:
        doc.add_heading(f'{week.week_label} ({week.date_range})', level=1)
        
        doc.add_heading('Alınan Çekler', level=2)
        for check in week.received_checks:
            doc.add_paragraph(f'• Çek No: {check.check_number} - ₺{check.amount:,.2f} - {check.due_date} - {check.bank_name}')
        
        doc.add_heading('Verilen Çekler', level=2)
        for check in week.issued_checks:
            doc.add_paragraph(f'• Çek No: {check.check_number} - ₺{check.amount:,.2f} - {check.due_date} - {check.bank_name}')
        
        doc.add_heading('Vadesi Gelen Faturalar', level=2)
        for invoice in week.invoices_due:
            doc.add_paragraph(f'• Fatura No: {invoice.invoice_number} - {invoice.customer_name} - ₺{invoice.amount:,.2f} - {invoice.due_date}')
        
        doc.add_paragraph('')
        doc.add_paragraph(f'Toplam Alacak: ₺{week.total_receivable:,.2f}')
        doc.add_paragraph(f'Toplam Borç: ₺{week.total_payable:,.2f}')
        doc.add_paragraph('')
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=haftalik_plan_{datetime.now().strftime('%Y%m%d')}.docx"}
    )

async def export_weekly_schedule_pdf(schedule):
    """Export weekly schedule to PDF format"""
    output = BytesIO()
    doc = SimpleDocTemplate(output, pagesize=landscape(A4))
    elements = []
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=20, textColor=colors.HexColor('#366092'))
    elements.append(Paragraph('Haftalık Ödeme Planı', title_style))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph(f'Rapor Tarihi: {datetime.now().strftime("%d.%m.%Y %H:%M")}', styles['Normal']))
    elements.append(Spacer(1, 20))
    
    for week in schedule:
        elements.append(Paragraph(f'{week.week_label} ({week.date_range})', styles['Heading2']))
        elements.append(Spacer(1, 10))
        
        # Received checks table
        if week.received_checks:
            elements.append(Paragraph('Alınan Çekler:', styles['Heading3']))
            data = [['Çek No', 'Tutar (₺)', 'Vade', 'Banka', 'Durum']]
            for check in week.received_checks:
                data.append([
                    check.check_number,
                    f"₺{check.amount:,.2f}",
                    check.due_date,
                    check.bank_name[:20],
                    check.status
                ])
            table = Table(data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#366092')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            elements.append(table)
            elements.append(Spacer(1, 10))
        
        # Issued checks table
        if week.issued_checks:
            elements.append(Paragraph('Verilen Çekler:', styles['Heading3']))
            data = [['Çek No', 'Tutar (₺)', 'Vade', 'Banka', 'Durum']]
            for check in week.issued_checks:
                data.append([
                    check.check_number,
                    f"₺{check.amount:,.2f}",
                    check.due_date,
                    check.bank_name[:20],
                    check.status
                ])
            table = Table(data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#366092')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            elements.append(table)
            elements.append(Spacer(1, 10))
        
        elements.append(Paragraph(f'Toplam Alacak: ₺{week.total_receivable:,.2f}', styles['Normal']))
        elements.append(Paragraph(f'Toplam Borç: ₺{week.total_payable:,.2f}', styles['Normal']))
        elements.append(Spacer(1, 20))
    
    doc.build(elements)
    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=haftalik_plan_{datetime.now().strftime('%Y%m%d')}.pdf"}
    )

# Export/Import endpoints
@api_router.get("/export/invoices")
async def export_invoices(format: str = "xlsx", user_id: str = Depends(get_current_user)):
    """Export all invoices to xlsx, docx, or pdf"""
    invoices = await db.invoices.find({}, {"_id": 0}).to_list(length=None)
    
    if format == "xlsx":
        return await export_invoices_xlsx(invoices)
    elif format == "docx":
        return await export_invoices_docx(invoices)
    elif format == "pdf":
        return await export_invoices_pdf(invoices)
    else:
        raise HTTPException(status_code=400, detail="Desteklenmeyen format")

@api_router.get("/export/checks")
async def export_checks(format: str = "xlsx", user_id: str = Depends(get_current_user)):
    """Export all checks to xlsx, docx, or pdf"""
    checks = await db.checks.find({}, {"_id": 0}).to_list(length=None)
    
    if format == "xlsx":
        return await export_checks_xlsx(checks)
    elif format == "docx":
        return await export_checks_docx(checks)
    elif format == "pdf":
        return await export_checks_pdf(checks)
    else:
        raise HTTPException(status_code=400, detail="Desteklenmeyen format")

@api_router.get("/export/payments")
async def export_payments(format: str = "xlsx", user_id: str = Depends(get_current_user)):
    """Export all payments to xlsx, docx, or pdf"""
    payments = await db.payments.find({}, {"_id": 0}).to_list(length=None)
    
    if format == "xlsx":
        return await export_payments_xlsx(payments)
    elif format == "docx":
        return await export_payments_docx(payments)
    elif format == "pdf":
        return await export_payments_pdf(payments)
    else:
        raise HTTPException(status_code=400, detail="Desteklenmeyen format")

@api_router.get("/export/weekly-schedule")
async def export_weekly_schedule(format: str = "xlsx", user_id: str = Depends(get_current_user)):
    """Export weekly schedule to xlsx, docx, or pdf"""
    # Get the schedule data
    schedule = await get_weekly_payment_schedule(weeks=4, user_id=user_id)
    
    if format == "xlsx":
        return await export_weekly_schedule_xlsx(schedule)
    elif format == "docx":
        return await export_weekly_schedule_docx(schedule)
    elif format == "pdf":
        return await export_weekly_schedule_pdf(schedule)
    else:
        raise HTTPException(status_code=400, detail="Desteklenmeyen format")

@api_router.post("/import/invoices")
async def import_invoices(file: UploadFile = File(...), user_id: str = Depends(get_current_user)):
    """Import invoices from xlsx file"""
    try:
        # Get user info for created_by
        user = await db.users.find_one({"id": user_id}, {"_id": 0})
        if not user:
            raise HTTPException(status_code=404, detail="Kullanıcı bulunamadı")
        
        # Read the Excel file
        contents = await file.read()
        df = pd.read_excel(BytesIO(contents))
        
        imported_count = 0
        for _, row in df.iterrows():
            invoice_data = {
                "id": str(uuid.uuid4()),
                "customer_id": str(row.get("customer_id", "")),
                "customer_name": str(row.get("customer_name", "")),
                "invoice_number": str(row.get("invoice_number", "")),
                "amount": float(row.get("amount", 0)),
                "paid_amount": float(row.get("paid_amount", 0)),
                "due_date": str(row.get("due_date", "")),
                "status": str(row.get("status", "unpaid")),
                "notes": str(row.get("notes", "")) if pd.notna(row.get("notes")) else None,
                "created_by": user_id,
                "created_by_username": user["username"],
                "created_at": datetime.now(timezone.utc).isoformat()
            }
            await db.invoices.insert_one(invoice_data)
            imported_count += 1
        
        return {"message": f"{imported_count} fatura başarıyla içe aktarıldı", "count": imported_count}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"İçe aktarma hatası: {str(e)}")

@api_router.post("/import/checks")
async def import_checks(file: UploadFile = File(...), user_id: str = Depends(get_current_user)):
    """Import checks from xlsx file"""
    try:
        # Get user info for created_by
        user = await db.users.find_one({"id": user_id}, {"_id": 0})
        if not user:
            raise HTTPException(status_code=404, detail="Kullanıcı bulunamadı")
        
        # Read the Excel file
        contents = await file.read()
        df = pd.read_excel(BytesIO(contents))
        
        imported_count = 0
        for _, row in df.iterrows():
            check_data = {
                "id": str(uuid.uuid4()),
                "check_type": str(row.get("check_type", "received")),
                "check_number": str(row.get("check_number", "")),
                "amount": float(row.get("amount", 0)),
                "due_date": str(row.get("due_date", "")),
                "bank_name": str(row.get("bank_name", "")),
                "payer_payee": str(row.get("payer_payee", "")),
                "status": str(row.get("status", "pending")),
                "notes": str(row.get("notes", "")) if pd.notna(row.get("notes")) else None,
                "created_by": user_id,
                "created_by_username": user["username"],
                "created_at": datetime.now(timezone.utc).isoformat()
            }
            await db.checks.insert_one(check_data)
            imported_count += 1
        
        return {"message": f"{imported_count} çek başarıyla içe aktarıldı", "count": imported_count}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"İçe aktarma hatası: {str(e)}")

@api_router.post("/import/payments")
async def import_payments(file: UploadFile = File(...), user_id: str = Depends(get_current_user)):
    """Import payments from xlsx file"""
    try:
        # Get user info for created_by
        user = await db.users.find_one({"id": user_id}, {"_id": 0})
        if not user:
            raise HTTPException(status_code=404, detail="Kullanıcı bulunamadı")
        
        # Read the Excel file
        contents = await file.read()
        df = pd.read_excel(BytesIO(contents))
        
        imported_count = 0
        for _, row in df.iterrows():
            payment_data = {
                "id": str(uuid.uuid4()),
                "invoice_id": str(row.get("invoice_id", "")),
                "invoice_number": str(row.get("invoice_number", "")) if pd.notna(row.get("invoice_number")) else None,
                "customer_name": str(row.get("customer_name", "")) if pd.notna(row.get("customer_name")) else None,
                "check_number": str(row.get("check_number", "")),
                "check_date": str(row.get("check_date", "")),
                "bank_name": str(row.get("bank_name", "")),
                "amount": float(row.get("amount", 0)),
                "created_by": user_id,
                "created_by_username": user["username"],
                "payment_date": datetime.now(timezone.utc).isoformat(),
                "created_at": datetime.now(timezone.utc).isoformat()
            }
            await db.payments.insert_one(payment_data)
            imported_count += 1
        
        return {"message": f"{imported_count} ödeme başarıyla içe aktarıldı", "count": imported_count}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"İçe aktarma hatası: {str(e)}")

# Dashboard export endpoint
@api_router.get("/export/dashboard-stats")
async def export_dashboard_stats(format: str = "xlsx", user_id: str = Depends(get_current_user)):
    """Export dashboard statistics summary"""
    try:
        # Get all statistics
        invoices = await db.invoices.find({}, {"_id": 0}).to_list(length=None)
        checks = await db.checks.find({}, {"_id": 0}).to_list(length=None)
        payments = await db.payments.find({}, {"_id": 0}).to_list(length=None)
        
        # Calculate statistics
        total_invoices = len(invoices)
        total_amount = sum(inv.get("amount", 0) for inv in invoices)
        paid_amount = sum(inv.get("paid_amount", 0) for inv in invoices)
        outstanding_amount = total_amount - paid_amount
        
        unpaid_count = len([inv for inv in invoices if inv.get("status") == "unpaid"])
        partial_count = len([inv for inv in invoices if inv.get("status") == "partial"])
        paid_count = len([inv for inv in invoices if inv.get("status") == "paid"])
        
        # Check statistics
        received_checks = [c for c in checks if c.get("check_type") == "received"]
        issued_checks = [c for c in checks if c.get("check_type") == "issued"]
        
        total_received_checks = len(received_checks)
        total_received_amount = sum(c.get("amount", 0) for c in received_checks)
        pending_received_checks = len([c for c in received_checks if c.get("status") == "pending"])
        
        total_issued_checks = len(issued_checks)
        total_issued_amount = sum(c.get("amount", 0) for c in issued_checks)
        pending_issued_checks = len([c for c in issued_checks if c.get("status") == "pending"])
        
        stats = {
            "total_invoices": total_invoices,
            "total_amount": total_amount,
            "paid_amount": paid_amount,
            "outstanding_amount": outstanding_amount,
            "unpaid_count": unpaid_count,
            "partial_count": partial_count,
            "paid_count": paid_count,
            "total_received_checks": total_received_checks,
            "total_received_amount": total_received_amount,
            "pending_received_checks": pending_received_checks,
            "total_issued_checks": total_issued_checks,
            "total_issued_amount": total_issued_amount,
            "pending_issued_checks": pending_issued_checks,
            "total_payments": len(payments)
        }
        
        if format == "xlsx":
            return await export_dashboard_stats_xlsx(stats)
        elif format == "docx":
            return await export_dashboard_stats_docx(stats)
        elif format == "pdf":
            return await export_dashboard_stats_pdf(stats)
        else:
            raise HTTPException(status_code=400, detail="Desteklenmeyen format")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Dışa aktarma hatası: {str(e)}")

async def export_dashboard_stats_xlsx(stats):
    """Export dashboard stats to Excel"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Dashboard İstatistikleri"
    
    # Title
    ws['A1'] = "Dashboard İstatistikleri - Özet Rapor"
    ws['A1'].font = Font(bold=True, size=16, color="FFFFFF")
    ws['A1'].fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    ws.merge_cells('A1:B1')
    
    ws['A2'] = f"Rapor Tarihi: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
    ws.append([])
    
    # Fatura İstatistikleri
    ws.append(["FATURA İSTATİSTİKLERİ", ""])
    ws[f'A{ws.max_row}'].font = Font(bold=True, size=12)
    ws.append(["Toplam Fatura Sayısı", stats["total_invoices"]])
    ws.append(["Toplam Tutar", f"₺{stats['total_amount']:,.2f}"])
    ws.append(["Ödenen Tutar", f"₺{stats['paid_amount']:,.2f}"])
    ws.append(["Kalan Tutar", f"₺{stats['outstanding_amount']:,.2f}"])
    ws.append(["Ödenmemiş Fatura", stats["unpaid_count"]])
    ws.append(["Kısmi Ödenmiş Fatura", stats["partial_count"]])
    ws.append(["Ödenmiş Fatura", stats["paid_count"]])
    ws.append([])
    
    # Çek İstatistikleri
    ws.append(["ÇEK İSTATİSTİKLERİ", ""])
    ws[f'A{ws.max_row}'].font = Font(bold=True, size=12)
    ws.append(["Alınan Çek Sayısı", stats["total_received_checks"]])
    ws.append(["Alınan Çek Tutarı", f"₺{stats['total_received_amount']:,.2f}"])
    ws.append(["Bekleyen Alınan Çek", stats["pending_received_checks"]])
    ws.append(["Verilen Çek Sayısı", stats["total_issued_checks"]])
    ws.append(["Verilen Çek Tutarı", f"₺{stats['total_issued_amount']:,.2f}"])
    ws.append(["Bekleyen Verilen Çek", stats["pending_issued_checks"]])
    ws.append([])
    
    # Ödeme İstatistikleri
    ws.append(["ÖDEME İSTATİSTİKLERİ", ""])
    ws[f'A{ws.max_row}'].font = Font(bold=True, size=12)
    ws.append(["Toplam Ödeme Sayısı", stats["total_payments"]])
    
    # Column widths
    ws.column_dimensions['A'].width = 30
    ws.column_dimensions['B'].width = 20
    
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename=dashboard_ozet_{datetime.now().strftime('%Y%m%d')}.xlsx"}
    )

async def export_dashboard_stats_docx(stats):
    """Export dashboard stats to Word"""
    doc = Document()
    doc.add_heading('Dashboard İstatistikleri - Özet Rapor', 0)
    doc.add_paragraph(f'Rapor Tarihi: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    doc.add_paragraph('')
    
    # Fatura İstatistikleri
    doc.add_heading('Fatura İstatistikleri', level=1)
    doc.add_paragraph(f'Toplam Fatura Sayısı: {stats["total_invoices"]}')
    doc.add_paragraph(f'Toplam Tutar: ₺{stats["total_amount"]:,.2f}')
    doc.add_paragraph(f'Ödenen Tutar: ₺{stats["paid_amount"]:,.2f}')
    doc.add_paragraph(f'Kalan Tutar: ₺{stats["outstanding_amount"]:,.2f}')
    doc.add_paragraph(f'Ödenmemiş Fatura: {stats["unpaid_count"]}')
    doc.add_paragraph(f'Kısmi Ödenmiş Fatura: {stats["partial_count"]}')
    doc.add_paragraph(f'Ödenmiş Fatura: {stats["paid_count"]}')
    doc.add_paragraph('')
    
    # Çek İstatistikleri
    doc.add_heading('Çek İstatistikleri', level=1)
    doc.add_paragraph(f'Alınan Çek Sayısı: {stats["total_received_checks"]}')
    doc.add_paragraph(f'Alınan Çek Tutarı: ₺{stats["total_received_amount"]:,.2f}')
    doc.add_paragraph(f'Bekleyen Alınan Çek: {stats["pending_received_checks"]}')
    doc.add_paragraph(f'Verilen Çek Sayısı: {stats["total_issued_checks"]}')
    doc.add_paragraph(f'Verilen Çek Tutarı: ₺{stats["total_issued_amount"]:,.2f}')
    doc.add_paragraph(f'Bekleyen Verilen Çek: {stats["pending_issued_checks"]}')
    doc.add_paragraph('')
    
    # Ödeme İstatistikleri
    doc.add_heading('Ödeme İstatistikleri', level=1)
    doc.add_paragraph(f'Toplam Ödeme Sayısı: {stats["total_payments"]}')
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=dashboard_ozet_{datetime.now().strftime('%Y%m%d')}.docx"}
    )

async def export_dashboard_stats_pdf(stats):
    """Export dashboard stats to PDF"""
    output = BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4)
    elements = []
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=20, textColor=colors.HexColor('#366092'))
    elements.append(Paragraph('Dashboard İstatistikleri - Özet Rapor', title_style))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph(f'Rapor Tarihi: {datetime.now().strftime("%d.%m.%Y %H:%M")}', styles['Normal']))
    elements.append(Spacer(1, 20))
    
    # Fatura İstatistikleri
    elements.append(Paragraph('Fatura İstatistikleri', styles['Heading2']))
    elements.append(Spacer(1, 10))
    
    invoice_data = [
        ['Metrik', 'Değer'],
        ['Toplam Fatura Sayısı', str(stats["total_invoices"])],
        ['Toplam Tutar', f"₺{stats['total_amount']:,.2f}"],
        ['Ödenen Tutar', f"₺{stats['paid_amount']:,.2f}"],
        ['Kalan Tutar', f"₺{stats['outstanding_amount']:,.2f}"],
        ['Ödenmemiş', str(stats["unpaid_count"])],
        ['Kısmi Ödenmiş', str(stats["partial_count"])],
        ['Ödenmiş', str(stats["paid_count"])]
    ]
    
    table = Table(invoice_data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#366092')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    elements.append(table)
    elements.append(Spacer(1, 20))
    
    # Çek İstatistikleri
    elements.append(Paragraph('Çek İstatistikleri', styles['Heading2']))
    elements.append(Spacer(1, 10))
    
    check_data = [
        ['Metrik', 'Değer'],
        ['Alınan Çek Sayısı', str(stats["total_received_checks"])],
        ['Alınan Çek Tutarı', f"₺{stats['total_received_amount']:,.2f}"],
        ['Bekleyen Alınan Çek', str(stats["pending_received_checks"])],
        ['Verilen Çek Sayısı', str(stats["total_issued_checks"])],
        ['Verilen Çek Tutarı', f"₺{stats['total_issued_amount']:,.2f}"],
        ['Bekleyen Verilen Çek', str(stats["pending_issued_checks"])]
    ]
    
    table = Table(check_data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#366092')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    elements.append(table)
    
    doc.build(elements)
    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=dashboard_ozet_{datetime.now().strftime('%Y%m%d')}.pdf"}
    )

# Logo management endpoints
@api_router.post("/settings/logo")
async def upload_logo(file: UploadFile = File(...), admin_id: str = Depends(get_current_admin_user)):
    """Upload company logo (admin only)"""
    try:
        # Validate file type
        if not file.content_type or not file.content_type.startswith('image/'):
            raise HTTPException(status_code=400, detail="Sadece resim dosyaları yüklenebilir")
        
        if not file.filename.lower().endswith('.png'):
            raise HTTPException(status_code=400, detail="Sadece PNG formatı desteklenmektedir")
        
        # Read file content
        contents = await file.read()
        
        # Save to database as base64
        logo_data = {
            "id": "company_logo",
            "filename": file.filename,
            "content_type": file.content_type,
            "data": base64.b64encode(contents).decode('utf-8'),
            "uploaded_by": admin_id,
            "uploaded_at": datetime.now(timezone.utc).isoformat()
        }
        
        # Upsert (update if exists, insert if not)
        await db.settings.update_one(
            {"id": "company_logo"},
            {"$set": logo_data},
            upsert=True
        )
        
        return {"message": "Logo başarıyla yüklendi", "filename": file.filename}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Logo yükleme hatası: {str(e)}")

@api_router.get("/settings/logo")
async def get_logo():
    """Get company logo (public endpoint)"""
    try:
        logo = await db.settings.find_one({"id": "company_logo"}, {"_id": 0})
        
        if not logo:
            raise HTTPException(status_code=404, detail="Logo bulunamadı")
        
        # Decode base64 and return image
        image_data = base64.b64decode(logo["data"])
        
        return StreamingResponse(
            BytesIO(image_data),
            media_type=logo.get("content_type", "image/png"),
            headers={"Content-Disposition": f'inline; filename="{logo.get("filename", "logo.png")}"'}
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Logo getirme hatası: {str(e)}")

@api_router.delete("/settings/logo")
async def delete_logo(admin_id: str = Depends(get_current_admin_user)):
    """Delete company logo (admin only)"""
    try:
        result = await db.settings.delete_one({"id": "company_logo"})
        
        if result.deleted_count == 0:
            raise HTTPException(status_code=404, detail="Logo bulunamadı")
        
        return {"message": "Logo başarıyla silindi"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Logo silme hatası: {str(e)}")


# ===== COMPANY INFO ENDPOINTS =====
@api_router.get("/company-info", response_model=CompanyInfo)
async def get_company_info(user_id: str = Depends(get_current_user)):
    """Get company information"""
    company = await db.company_info.find_one({}, {"_id": 0})
    if not company:
        # Return default empty company info
        return CompanyInfo(
            company_name="",
            bank_accounts=[]
        )
    return CompanyInfo(**company)

@api_router.post("/company-info", response_model=CompanyInfo)
async def create_or_update_company_info(
    company_data: CompanyInfoCreate,
    user_id: str = Depends(get_current_user)
):
    """Create or update company information (admin only)"""
    # Check if user is admin
    user = await db.users.find_one({"id": user_id}, {"_id": 0})
    if not user or not user.get("is_admin"):
        raise HTTPException(status_code=403, detail="Admin yetkilisi gerekli")
    
    # Check if company info exists
    existing = await db.company_info.find_one({})
    
    company_obj = CompanyInfo(**company_data.model_dump())
    company_obj.updated_at = datetime.now(timezone.utc).isoformat()
    
    if existing:
        # Update existing
        company_obj.id = existing.get("id", company_obj.id)
        company_obj.created_at = existing.get("created_at", company_obj.created_at)
        await db.company_info.replace_one({"id": company_obj.id}, company_obj.model_dump())
    else:
        # Create new
        await db.company_info.insert_one(company_obj.model_dump())
    
    return company_obj

@api_router.get("/company-info/banks")
async def get_bank_accounts(user_id: str = Depends(get_current_user)):
    """Get list of company bank accounts"""
    company = await db.company_info.find_one({}, {"_id": 0})
    if not company:
        return []
    
    company_info = CompanyInfo(**company)
    return company_info.bank_accounts


app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("startup")
async def startup_event():
    """Start the scheduler on application startup"""
    # Schedule daily check at 12:00 PM Turkish time
    scheduler.add_job(
        check_upcoming_invoices,
        CronTrigger(hour=12, minute=0, timezone=pytz.timezone('Europe/Istanbul')),
        id='invoice_reminder_check',
        name='Daily Invoice Reminder Check',
        replace_existing=True
    )
    scheduler.start()
    logger.info("Scheduler started - Daily invoice reminders will be sent at 12:00 PM Turkish time")

@app.on_event("shutdown")
async def shutdown_db_client():
    scheduler.shutdown()
    client.close()
    logger.info("Scheduler and database connection closed")
